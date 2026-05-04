from __future__ import annotations

import os
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(os.environ["THESIS_SRC"])
DST = Path(os.environ["THESIS_DST"])


CN_ABSTRACT = (
    "随着高校后勤服务数字化以及移动端应用场景的持续扩展，校园餐饮服务逐渐从窗口排队、电话订餐等线下方式，"
    "转向小程序点餐、商户在线管理以及订单状态实时跟踪。针对校园外卖点餐过程中存在的商户信息分散、菜品维护不及时、"
    "订单流转依赖人工沟通以及运营数据难以沉淀等问题，本课题围绕校园场景设计并实现了一套餐饮外卖点餐系统。"
    "系统选用 Spring Boot 构建后端服务，运用 MyBatis 进行数据访问，以 MySQL 存储用户、商户、菜品、购物车和订单等核心数据，"
    "并借助 Redis、JWT、WebSocket 等组件支撑状态辅助、身份认证以及订单提醒。小程序端基于 uni-app 开发，Web 管理端以 Nginx "
    "静态资源形式提供商户和运营管理入口。\n"
    "论文在需求建模基础上，对系统架构、数据库模型、关键业务流程、接口边界和质量验证进行了分析。系统以共享业务表中的 merchant_id、"
    "campus_id 以及 account_type 等字段表达商户作用域，并在登录认证、商户切换、购物车隔离、订单提交、订单通知和运营统计等流程中传递商户上下文。"
    "验证部分围绕登录认证、商品浏览、购物车隔离、订单生成、管理端处理、WebSocket 通知和统计报表等路径展开，重点检查功能闭环、权限边界和数据一致性。"
    "结果表明，系统能够覆盖校园外卖点餐的基本业务链路，并为多商户校园餐饮服务提供了可运行的工程基础。对于真实支付生产上线、大规模压力测试、"
    "长期稳定性和正式用户调研等缺乏充分证据的内容，论文只作为边界条件和后续改进方向进行说明。"
)

EN_ABSTRACT = (
    "With the continuous development of digital campus logistics services and mobile application scenarios, campus catering services are gradually shifting "
    "from offline modes such as window queuing and telephone ordering to mini-program ordering, online merchant management, and real-time order status tracking. "
    "To address problems such as scattered merchant information, delayed dish maintenance, manual communication in order processing, and insufficient accumulation "
    "of operational data, this project designs and implements a campus-oriented food delivery ordering system. The back-end service is built with Spring Boot and "
    "MyBatis, while MySQL stores core business data such as users, merchants, dishes, shopping carts, and orders. Redis, JWT, and WebSocket are used to support "
    "state assistance, identity authentication, and order notification. The mobile client is developed with uni-app, and the Web management terminal is provided "
    "as static resources deployed through Nginx.\n"
    "Based on requirement modeling, the thesis analyzes system architecture, database design, key business processes, interface boundaries, and quality verification. "
    "The system expresses merchant scope through fields such as merchant_id, campus_id, and account_type in shared business tables, and transfers merchant context "
    "across login authentication, merchant switching, shopping cart isolation, order submission, order notification, and operational statistics. The verification work "
    "focuses on login authentication, dish browsing, shopping cart isolation, order generation, management-side processing, WebSocket notification, and report statistics, "
    "with particular attention to functional closure, permission boundaries, and data consistency. The results show that the system covers the basic business loop of "
    "campus food delivery ordering and provides an engineering foundation for multi-merchant campus catering services. Capabilities without sufficient evidence, such as "
    "production-level payment deployment, large-scale stress testing, long-term stability, and formal user research, are treated only as boundary conditions and future "
    "improvement directions."
)


TECH_TABLE = [
    ["技术点", "选用方案", "可替代方案", "选用缘由", "当前边界"],
    ["后端架构", "Spring Boot 分层单体", "Spring Cloud 微服务", "校园商户规模可控，单体部署复杂度较低，控制器、服务层和 Mapper 层边界清晰", "不具备独立服务拆分和服务治理能力"],
    ["数据访问", "MyBatis", "JPA、MyBatis-Plus", "SQL 可控性较强，适宜订单、报表和多条件查询", "需要持续检查 SQL 中的商户过滤条件"],
    ["小程序端", "uni-app", "原生微信小程序", "可复用 Vue 语法和 uni-app 页面组织方式，适宜移动端点餐场景", "当前重点是微信小程序端，不展开多端适配验证"],
    ["身份认证", "JWT", "Session、OAuth2", "适合前后端分离接口中的无状态身份传递", "未包含刷新 token、黑名单和完整身份治理"],
    ["实时通知", "WebSocket", "SSE、轮询、消息队列", "适合订单提醒和催单这类实时消息", "未形成消息确认、离线补偿和失败重试机制"],
    ["商户隔离", "共享表 + merchant_id", "独立数据库、独立 schema、数据库 RLS", "实现成本较低，适合校园内有限商户数量", "属于行级作用域控制，不是严格物理隔离"],
]

ARCH_TABLE = [
    ["层次", "组成", "主要职责", "交互方式"],
    ["用户端", "uni-app 微信小程序", "商户浏览、菜品选择、购物车维护、订单提交和历史订单查看", "访问 /user 前缀接口，请求头携带 authentication"],
    ["管理端", "Nginx 静态 Web 管理端", "商户维护、商品维护、订单处理、工作台和报表查看", "访问 /admin 前缀接口，请求头携带 token"],
    ["后端服务", "Spring Boot + MyBatis", "认证、权限、商户作用域、订单状态流转和统计逻辑", "提供 REST 接口并访问 Mapper 层"],
    ["数据与状态", "MySQL、Redis", "持久化业务数据，并对缓存或状态处理提供辅助", "后端服务统一访问"],
    ["通知组件", "WebSocketServer /ws/{sid}", "发送新订单和催单等订单事件", "向商户员工或平台管理员推送消息"],
]

DEPLOY_TABLE = [
    ["部署对象", "位置或形式", "说明"],
    ["Web 管理端资源", "core/nginx/html/sky", "以静态资源形式由 Nginx 提供访问入口"],
    ["后端服务", "sky-server Spring Boot 应用", "承接 /user 与 /admin 接口，并连接数据库和缓存"],
    ["数据库", "MySQL", "保存用户、商户、商品、购物车、订单和订单明细等数据"],
    ["缓存或状态辅助", "Redis", "用于缓存或状态辅助，论文不扩大为完整缓存治理体系"],
    ["外部接口", "对象存储、微信登录或支付", "按配置接入，生产级支付能力只作为边界说明"],
]

ER_TABLE = [
    ["实体关系", "含义"],
    ["campus 1:N merchant", "一个校区可以包含多个校园商户"],
    ["merchant 1:N employee", "商户与后台员工通过 merchant_id 建立管理范围"],
    ["merchant 1:N category/dish/setmeal", "分类、菜品和套餐通过 merchant_id 表达所属商户"],
    ["user 1:N address_book", "学生用户可以维护多个收货地址"],
    ["user + merchant 1:N shopping_cart", "购物车同时按用户和商户隔离"],
    ["user + merchant + campus 1:N orders", "订单记录用户、商户和校区归属"],
    ["orders 1:N order_detail", "订单明细保存具体商品、数量和金额"],
]

ROLE_TABLE = [
    ["角色", "主要入口", "身份依据", "主要权限", "商户边界"],
    ["学生用户", "微信小程序", "user token 与 user_id", "浏览商户与菜品、维护购物车、提交订单、查看历史订单、催单", "只能访问本人订单和购物车"],
    ["商户员工", "Web 管理端", "token、account_type、merchant_id", "处理所属商户订单、查看所属商户业务数据", "由 merchant_id 限定"],
    ["商户管理员", "Web 管理端", "token、account_type、merchant_id", "管理所属商户商品、订单和基础运营信息", "由 merchant_id 限定"],
    ["平台管理员", "Web 管理端", "token、account_type", "管理商户、查看平台或指定商户数据", "可按权限查看更大范围"],
]

TEST_TABLE = [
    ["编号", "测试对象", "前置条件", "操作步骤", "输入数据", "预期结果", "实际记录方式", "结论"],
    ["TC-01", "登录认证", "后端服务正常，存在用户或员工账号", "登录后携带 token 访问受保护接口，再移除 token 访问同一接口", "authentication 或 token", "有效 token 可访问，无效或缺失 token 被拒绝", "接口响应、后端日志", "待执行记录"],
    ["TC-02", "商户切换", "至少存在两个启用商户", "小程序端切换商户后刷新分类和菜品", "merchantId/shopId", "页面只展示当前商户商品", "页面截图、接口返回", "待执行记录"],
    ["TC-03", "购物车隔离", "同一用户可访问两个商户", "分别向商户 A、商户 B 加入商品，再分别查询购物车", "user_id、merchant_id", "两个商户购物车互不影响", "接口返回、数据库记录", "待执行记录"],
    ["TC-04", "订单提交", "当前商户购物车存在商品，地址有效", "提交订单并查看订单详情", "addressBookId、merchantId", "生成订单主表和明细，清理当前商户购物车", "接口返回、数据库记录", "待执行记录"],
    ["TC-05", "空购物车下单", "当前商户购物车为空", "直接提交订单", "addressBookId、merchantId", "返回业务异常，不生成订单", "接口返回、数据库记录", "待执行记录"],
    ["TC-06", "管理端订单处理", "存在待确认订单", "商户端执行确认、派送、完成操作", "orderId", "订单状态按流程变更", "接口返回、订单状态", "待执行记录"],
    ["TC-07", "权限边界", "存在商户账号和其他商户订单", "商户账号查询非所属商户订单", "token、merchantId", "拒绝越权访问或返回空范围数据", "接口响应、日志", "待执行记录"],
    ["TC-08", "订单通知", "WebSocket 连接建立", "触发支付成功或催单流程", "orderNumber/orderId", "商户端收到 NEW_ORDER 或 ORDER_REMINDER 消息", "WebSocket 消息记录", "待执行记录"],
    ["TC-09", "统计报表", "存在已完成订单数据", "查询营业额、订单数量、销量排行", "begin、end、merchantId", "返回指定时间和商户范围内统计数据", "接口返回", "待执行记录"],
    ["TC-10", "支付边界", "系统处于模拟或测试支付配置", "提交订单后执行支付接口", "orderNumber", "可验证订单状态流转，但不声明生产结算能力", "接口返回、状态记录", "待执行记录"],
]


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    for idx, part in enumerate(text.split("\n")):
        if idx:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.font.size = Pt(12)


def find_paragraph(doc: Document, exact: str, start: int = 0):
    for i, paragraph in enumerate(doc.paragraphs[start:], start):
        if paragraph.text.strip() == exact:
            return i, paragraph
    raise ValueError(f"paragraph not found: {exact}")


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if style:
        new_paragraph.style = style
    if text:
        set_paragraph_text(new_paragraph, text)
    return new_paragraph


def insert_table_after(paragraph, rows: list[list[str]], style_name: str = "Table Grid"):
    doc = paragraph.part.document
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    try:
        table.style = style_name
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    if r == 0:
                        run.bold = True
    tbl = table._tbl
    paragraph._p.addnext(tbl)
    return table


def insert_captioned_table_after(paragraph, caption: str, rows: list[list[str]]):
    cap = insert_paragraph_after(paragraph, caption, "Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.bold = True
    table = insert_table_after(cap, rows)
    spacer = insert_paragraph_after(cap, "")
    table._tbl.addnext(spacer._p)
    return spacer


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    DST.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SRC, DST)

    doc = Document(DST)
    set_update_fields_on_open(doc)

    cn_index, _ = find_paragraph(doc, "摘    要")
    set_paragraph_text(doc.paragraphs[cn_index + 1], CN_ABSTRACT)
    set_paragraph_text(doc.paragraphs[cn_index + 2], "关键词：校园外卖点餐系统；多商户管理；微信小程序；Spring Boot；订单管理")

    en_index, _ = find_paragraph(doc, "ABSTRACT")
    set_paragraph_text(doc.paragraphs[en_index + 1], EN_ABSTRACT)
    set_paragraph_text(doc.paragraphs[en_index + 2], "Keywords: campus food delivery ordering system; multi-merchant management; WeChat mini program; Spring Boot; order management")

    _, p24 = find_paragraph(doc, "2.4 技术选型与适配性分析")
    after = insert_captioned_table_after(p24, "表2.1 系统关键技术选型对比", TECH_TABLE)
    insert_paragraph_after(after, "本系统在技术选型上更重视校园场景下的可落地性，而不是追求复杂架构的堆叠。Spring Boot 分层单体结构可以把控制器、业务服务、数据访问和配置管理组织在同一应用中，部署和调试成本相对较低。对于毕业设计阶段的校园外卖业务来说，商户数量、用户范围和业务规模通常处在可控范围内，因此没有必要一开始就引入完整的微服务拆分、服务注册、链路追踪和分布式事务治理。", "Normal")

    _, p41 = find_paragraph(doc, "4.1 总体架构设计")
    after = insert_captioned_table_after(p41, "表4.1 系统总体架构组成", ARCH_TABLE)
    after = insert_captioned_table_after(after, "表4.2 系统部署结构说明", DEPLOY_TABLE)

    _, p44 = find_paragraph(doc, "4.4 数据库实体与多商户隔离")
    after = insert_captioned_table_after(p44, "表4.3 核心实体关系说明", ER_TABLE)
    insert_paragraph_after(after, "需要说明的是，本系统的多商户能力不是物理意义上的独立数据库或独立 schema 隔离，而是在共享业务表中借助 merchant_id、campus_id 和 account_type 等字段进行行级作用域控制。迁移脚本为 employee、category、dish、setmeal、shopping_cart、orders 等表补充商户相关字段，并增加与商户、状态和时间相关的索引。运行时，MultiMerchantSchemaSupport 会检查多商户字段是否存在；服务层再借助 MerchantScopeUtils 解析当前账号能够访问的商户范围。", "Normal")

    _, p45 = find_paragraph(doc, "4.5 接口边界与安全设计")
    insert_captioned_table_after(p45, "表4.4 角色权限矩阵", ROLE_TABLE)

    _, p62 = find_paragraph(doc, "6.2 功能测试")
    insert_captioned_table_after(p62, "表6.2 系统主要测试用例", TEST_TABLE)

    _, appendix_b = find_paragraph(doc, "附录 B 主要接口与页面映射")
    p = insert_paragraph_after(appendix_b, "附录表B.2 典型接口请求与响应示例", "Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
    examples = [
        "用户端订单提交接口：POST /user/order/submit。请求体可包含 addressBookId、merchantId、shopId、remark、tablewareNumber 和 tablewareStatus。该接口需要用户端携带 authentication，用于生成订单主表和订单明细。",
        "用户端购物车查询接口：GET /user/shoppingCart/list?merchantId=1。该接口返回当前用户在指定商户下的购物车数据，数据库支持商户字段时需要同时依据 user_id 和 merchant_id 过滤。",
        "管理端订单查询接口：GET /admin/order/conditionSearch。常用参数包括 page、pageSize、status 和 merchantId。商户账号查询时，服务层应将查询范围限制在当前账号绑定的 merchant_id。",
        "管理端统计接口：GET /admin/report/turnoverStatistics、/admin/report/ordersStatistics、/admin/report/top10。统计接口需要明确 begin、end 和商户范围，避免混淆平台口径和商户口径。",
    ]
    cursor = p
    for example in examples:
        cursor = insert_paragraph_after(cursor, example, "Normal")

    doc.save(DST)
    print(DST)


if __name__ == "__main__":
    main()
