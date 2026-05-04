from __future__ import annotations

import hashlib
import json
import os
import zipfile
from pathlib import Path

from docx import Document


ORIGINAL = Path(os.environ["THESIS_SRC"])
UPGRADED = Path(os.environ["THESIS_DST"])
REPORT = Path(r"D:\sky-delivery\docs\thesis-upgrade\work\qa-report.json")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest().upper()


def doc_summary(path: Path) -> dict:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    rel_targets = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            rel_targets.append(str(rel.target_ref))
    header_texts = []
    footer_texts = []
    for section in doc.sections:
        header_texts.append("\n".join(p.text for p in section.header.paragraphs if p.text.strip()))
        footer_texts.append("\n".join(p.text for p in section.footer.paragraphs if p.text.strip()))
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "headings": headings,
        "image_relationships": rel_targets,
        "header_nonempty_count": sum(1 for text in header_texts if text.strip()),
        "footer_nonempty_count": sum(1 for text in footer_texts if text.strip()),
        "has_cn_abstract_upgrade": "共享业务表中的 merchant_id" in text,
        "has_tech_table_caption": "表2.1 系统关键技术选型对比" in text,
        "has_arch_table_caption": "表4.1 系统总体架构组成" in text,
        "has_er_table_caption": "表4.3 核心实体关系说明" in text,
        "has_role_table_caption": "表4.4 角色权限矩阵" in text,
        "has_test_table_caption": "表6.2 系统主要测试用例" in text,
        "has_appendix_api_examples": "附录表B.2 典型接口请求与响应示例" in text,
        "forbidden_phrase_hits": [
            phrase
            for phrase in ["本次修复", "Feynman 发现 bug 后已修复", "至于"]
            if phrase in text
        ],
    }


def zip_summary(path: Path) -> dict:
    required = {
        "[Content_Types].xml",
        "word/document.xml",
        "word/styles.xml",
        "word/settings.xml",
    }
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        corrupt = zf.testzip()
    return {
        "required_parts_present": sorted(required.intersection(names)),
        "missing_required_parts": sorted(required - names),
        "corrupt_member": corrupt,
        "part_count": len(names),
    }


def main() -> None:
    report = {
        "original_path": str(ORIGINAL),
        "upgraded_path": str(UPGRADED),
        "original_exists": ORIGINAL.exists(),
        "upgraded_exists": UPGRADED.exists(),
        "original_sha256": sha256(ORIGINAL),
        "upgraded_sha256": sha256(UPGRADED),
        "original_docx": doc_summary(ORIGINAL),
        "upgraded_docx": doc_summary(UPGRADED),
        "upgraded_zip": zip_summary(UPGRADED),
        "qa_mode": "non-render structural/openability checks only",
    }
    report["checks"] = {
        "upgraded_opens": report["upgraded_docx"]["paragraphs"] > 0,
        "zip_has_required_parts": not report["upgraded_zip"]["missing_required_parts"],
        "zip_not_corrupt": report["upgraded_zip"]["corrupt_member"] is None,
        "tables_added": report["upgraded_docx"]["tables"] > report["original_docx"]["tables"],
        "expected_content_present": all(
            report["upgraded_docx"][key]
            for key in [
                "has_cn_abstract_upgrade",
                "has_tech_table_caption",
                "has_arch_table_caption",
                "has_er_table_caption",
                "has_role_table_caption",
                "has_test_table_caption",
                "has_appendix_api_examples",
            ]
        ),
        "no_forbidden_phrases": not report["upgraded_docx"]["forbidden_phrase_hits"],
    }
    report["passes"] = all(report["checks"].values())
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"report": str(REPORT), "passes": report["passes"], "checks": report["checks"]}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
